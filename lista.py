import streamlit as st
import pandas as pd
import datetime
import os
import sqlite3
from io import BytesIO
from docx import Document

# Configuraci√≥n de la contrase√±a
PASSWORD = "defvm11"

# Estado de autenticaci√≥n
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

# P√°gina de inicio de sesi√≥n
if not st.session_state.authenticated:
    st.title("Sistema de Registro de Asistencia")
    password_input = st.text_input("Ingrese la contrase√±a:", type="password")
    if st.button("Ingresar"):
        if password_input == PASSWORD:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("Contrase√±a incorrecta")
    st.stop()

# Directorios de almacenamiento
ONEDRIVE_PATH = r"C:\\Users\\sup11\\OneDrive\\Attachments\\Documentos\\Interfaces de phyton\\Lista de asistencia"
DB_PATH = os.path.join(ONEDRIVE_PATH, "asistencia.db")
DOCX_PATH = os.path.join(ONEDRIVE_PATH, "lista.docx")
if not os.path.exists(ONEDRIVE_PATH):
    os.makedirs(ONEDRIVE_PATH)

# Conectar con SQLite y verificar la existencia de la base de datos
def conectar_db():
    if not os.path.exists(DB_PATH):
        st.error(f"La base de datos no existe en la ruta: {DB_PATH}")
        st.stop()
    try:
        conn = sqlite3.connect(DB_PATH)
        return conn
    except sqlite3.Error as e:
        st.error(f"Error al conectar con la base de datos: {e}")
        st.stop()

# Obtener lista de docentes
def obtener_docentes():
    try:
        conn = conectar_db()
        cursor = conn.cursor()
        cursor.execute("SELECT id, apellido_paterno, apellido_materno, nombre FROM docentes WHERE activo = 1")
        docentes = cursor.fetchall()
        conn.close()
        return [f"{d[1]} {d[2]} {d[3]}" for d in docentes]
    except sqlite3.OperationalError:
        st.error("Error al leer la tabla 'docentes'. Verifique que la base de datos est√© correctamente configurada.")
        return []

# Registro de actividad y fecha
st.title("Registro de Actividad")
actividad = st.text_input("Ingrese el nombre de la actividad:")
fecha_actividad = st.date_input("Seleccione la fecha de la actividad:")
if not actividad or not fecha_actividad:
    st.warning("Debe ingresar la actividad y la fecha antes de continuar.")
    st.stop()

# Archivo de asistencia
nombre_archivo = f"Asistencia_{actividad}_{fecha_actividad}.xlsx"
archivo_ruta = os.path.join(ONEDRIVE_PATH, nombre_archivo)
columnas = ["No.", "Nombre Completo", "Hora de Entrada", "Firma", "Hora de Salida", "Firma"]

# Verificar si el archivo Excel existe
if os.path.exists(archivo_ruta):
    try:
        df_asistencia = pd.read_excel(archivo_ruta, engine='openpyxl')
        
        # **üõ† Eliminar columnas duplicadas antes de concatenar**
        df_asistencia = df_asistencia.loc[:, ~df_asistencia.columns.duplicated()].copy()

        # **üõ† Verificar que el √≠ndice no tenga valores repetidos**
        df_asistencia = df_asistencia.reset_index(drop=True)

    except Exception as e:
        st.error(f"Error al leer el archivo de asistencia: {e}")
        df_asistencia = pd.DataFrame(columns=columnas)
else:
    df_asistencia = pd.DataFrame(columns=columnas)

# Formulario de registro de asistencia
st.title("Registro de Asistencia")
docentes = obtener_docentes()
docentes_seleccionados = st.multiselect("Seleccione los Docentes:", docentes)

hora_entrada_base = st.time_input("Hora de Entrada:")
hora_salida = st.time_input("Hora de Salida:")

if st.button("Registrar Asistencia"):
    if docentes_seleccionados:
        registros = []
        for i, docente in enumerate(docentes_seleccionados):
            hora_entrada = (datetime.datetime.combine(datetime.date.today(), hora_entrada_base) + datetime.timedelta(minutes=i)).time()
            registros.append([len(df_asistencia) + 1 + i, docente, str(hora_entrada), "", str(hora_salida), ""])

        df_nuevos = pd.DataFrame(registros, columns=columnas)

        # **üîß Validar que las columnas son correctas**
        df_nuevos = df_nuevos.loc[:, ~df_nuevos.columns.duplicated()].copy()

        df_asistencia = pd.concat([df_asistencia, df_nuevos], ignore_index=True)
        df_asistencia.to_excel(archivo_ruta, index=False, engine='openpyxl')

        st.success("Asistencia registrada correctamente")
    else:
        st.warning("Debe seleccionar al menos un docente.")

# Mostrar tabla de asistencia
st.subheader("Lista de Asistencia del d√≠a")
if not df_asistencia.empty:
    st.dataframe(df_asistencia)
else:
    st.warning("No hay registros de asistencia disponibles.")

# Generar documento en Word
def generar_docx():
    doc = Document()
    doc.add_paragraph("Servicios Educativos Integrados Al Estado de M√©xico\nSupervisi√≥n 11 de Educaci√≥n F√≠sica Valle de M√©xico")
    doc.add_paragraph(f"Lista de asistencia: {actividad}")
    doc.add_paragraph(f"Fecha: {fecha_actividad}")
    tabla = doc.add_table(rows=1, cols=len(columnas))
    for i, columna in enumerate(columnas):
        tabla.cell(0, i).text = columna
    for index, row in df_asistencia.iterrows():
        fila = tabla.add_row().cells
        for i, valor in enumerate(row):
            fila[i].text = str(valor)
    doc.add_paragraph("\nATENTAMENTE\nDOCTOR\nGUZMAN HERNANDEZ ESTRADA\nINSPECTOR DE LA SUPERVISI√ìN 11")
    doc.save(DOCX_PATH)
    return DOCX_PATH

# Bot√≥n para generar documento Word
if st.button("Generar Lista de Asistencia para Firma"):
    df_asistencia.to_excel(archivo_ruta, index=False, engine='openpyxl')
    st.success(f"Lista de asistencia guardada en: {archivo_ruta}")
    docx_path = generar_docx()
    with open(docx_path, "rb") as f:
        st.download_button("Descargar Lista de Asistencia en Word", f, file_name="Lista_Asistencia.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

